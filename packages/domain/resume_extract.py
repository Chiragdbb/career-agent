"""Text extraction from resume uploads (PDF / DOCX)."""

from __future__ import annotations

import io

from packages.domain.exceptions import DomainError

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

ALLOWED_MIME_TYPES = frozenset({PDF_MIME, DOCX_MIME})
ALLOWED_EXTENSIONS = frozenset({".pdf", ".docx"})


class UnsupportedResumeFormatError(DomainError):
    """Uploaded file is not a supported resume format."""


def detect_mime_type(filename: str, content_type: str | None) -> str:
    """Resolve MIME type from Content-Type header and/or filename extension."""
    normalized = (content_type or "").split(";")[0].strip().lower()
    if normalized in ALLOWED_MIME_TYPES:
        return normalized

    lower = filename.lower()
    if lower.endswith(".pdf"):
        return PDF_MIME
    if lower.endswith(".docx"):
        return DOCX_MIME

    raise UnsupportedResumeFormatError(
        "Only PDF and DOCX resumes are supported"
    )


def extract_text(data: bytes, mime_type: str) -> str:
    """Extract plain text from PDF or DOCX bytes. Does not invent content."""
    if mime_type == PDF_MIME:
        return _extract_pdf(data)
    if mime_type == DOCX_MIME:
        return _extract_docx(data)
    raise UnsupportedResumeFormatError(
        "Only PDF and DOCX resumes are supported"
    )


def _extract_pdf(data: bytes) -> str:
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:  # pragma: no cover
        raise DomainError("PyMuPDF is required for PDF extraction") from exc

    try:
        doc = fitz.open(stream=data, filetype="pdf")
    except Exception as exc:
        raise DomainError("Could not read PDF resume") from exc

    try:
        parts: list[str] = []
        for page in doc:
            text = page.get_text("text") or ""
            if text.strip():
                parts.append(text)
        return "\n".join(parts).strip()
    finally:
        doc.close()


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:  # pragma: no cover
        raise DomainError("python-docx is required for DOCX extraction") from exc

    try:
        document = DocxDocument(io.BytesIO(data))
    except Exception as exc:
        raise DomainError("Could not read DOCX resume") from exc

    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = (paragraph.text or "").strip()
        if text:
            parts.append(text)

    # Tables often hold contact / skills grids — extract cell text without inventing.
    for table in document.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()
